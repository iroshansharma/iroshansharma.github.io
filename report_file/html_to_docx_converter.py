#!/usr/bin/env python3
"""
HTML to DOCX Converter for CRM System Documentation
Converts multiple HTML files to a single, professionally formatted Word document
"""

import os
import re
from pathlib import Path
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

class HTMLToDocxConverter:
    def __init__(self):
        self.doc = Document()
        self.setup_styles()
        
    def setup_styles(self):
        """Setup custom styles for the document"""
        styles = self.doc.styles
        
        # Title style
        title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_font = title_style.font
        title_font.name = 'Segoe UI'
        title_font.size = Pt(28)
        title_font.bold = True
        title_font.color.rgb = RGBColor(0, 123, 255)  # Blue color
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(24)
        
        # Subtitle style
        subtitle_style = styles.add_style('CustomSubtitle', WD_STYLE_TYPE.PARAGRAPH)
        subtitle_font = subtitle_style.font
        subtitle_font.name = 'Segoe UI'
        subtitle_font.size = Pt(16)
        subtitle_font.color.rgb = RGBColor(108, 117, 125)  # Gray color
        subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_style.paragraph_format.space_after = Pt(12)
        
        # Chapter title style
        chapter_style = styles.add_style('ChapterTitle', WD_STYLE_TYPE.PARAGRAPH)
        chapter_font = chapter_style.font
        chapter_font.name = 'Segoe UI'
        chapter_font.size = Pt(24)
        chapter_font.bold = True
        chapter_font.color.rgb = RGBColor(40, 167, 69)  # Green color
        chapter_style.paragraph_format.space_before = Pt(24)
        chapter_style.paragraph_format.space_after = Pt(18)
        
        # Section heading styles (H2, H3, H4)
        h2_style = styles.add_style('CustomH2', WD_STYLE_TYPE.PARAGRAPH)
        h2_font = h2_style.font
        h2_font.name = 'Segoe UI'
        h2_font.size = Pt(18)
        h2_font.bold = True
        h2_font.color.rgb = RGBColor(40, 167, 69)
        h2_style.paragraph_format.space_before = Pt(18)
        h2_style.paragraph_format.space_after = Pt(12)
        
        h3_style = styles.add_style('CustomH3', WD_STYLE_TYPE.PARAGRAPH)
        h3_font = h3_style.font
        h3_font.name = 'Segoe UI'
        h3_font.size = Pt(16)
        h3_font.bold = True
        h3_font.color.rgb = RGBColor(32, 201, 151)
        h3_style.paragraph_format.space_before = Pt(14)
        h3_style.paragraph_format.space_after = Pt(10)
        
        h4_style = styles.add_style('CustomH4', WD_STYLE_TYPE.PARAGRAPH)
        h4_font = h4_style.font
        h4_font.name = 'Segoe UI'
        h4_font.size = Pt(14)
        h4_font.bold = True
        h4_font.color.rgb = RGBColor(85, 85, 85)
        h4_style.paragraph_format.space_before = Pt(12)
        h4_style.paragraph_format.space_after = Pt(8)
        
        # Code block style
        code_style = styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
        code_font = code_style.font
        code_font.name = 'Courier New'
        code_font.size = Pt(10)
        code_style.paragraph_format.left_indent = Inches(0.5)
        code_style.paragraph_format.space_before = Pt(6)
        code_style.paragraph_format.space_after = Pt(6)
        
        # Normal text style
        normal_style = styles['Normal']
        normal_font = normal_style.font
        normal_font.name = 'Segoe UI'
        normal_font.size = Pt(11)
        normal_style.paragraph_format.space_after = Pt(6)
        normal_style.paragraph_format.line_spacing = 1.15
        
    def add_title_page(self):
        """Add a professional title page"""
        # Main title
        title = self.doc.add_paragraph('Customer Relationship Management System', style='CustomTitle')
        
        # Subtitle
        subtitle = self.doc.add_paragraph('Comprehensive Technical Documentation Report', style='CustomSubtitle')
        
        # Add some space
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        # Project details
        details = self.doc.add_paragraph()
        details.alignment = WD_ALIGN_PARAGRAPH.CENTER
        details.add_run('A Complete 200-Page Analysis of CRM Implementation\n\n').font.size = Pt(14)
        details.add_run('Technology Stack: PHP, MySQL, HTML5, CSS3, JavaScript\n').font.size = Pt(12)
        details.add_run('Architecture: MVC Pattern with Role-based Access Control\n').font.size = Pt(12)
        details.add_run('Features: 15+ Comprehensive Modules\n').font.size = Pt(12)
        
        # Add page break
        self.doc.add_page_break()
        
    def process_html_content(self, html_content, chapter_num=None):
        """Process HTML content and convert to Word format"""
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Find the main content container
        container = soup.find('div', class_='container')
        if not container:
            container = soup.find('body')
        
        if not container:
            return
            
        # Process each element in the container
        for element in container.find_all(['h1', 'h2', 'h3', 'h4', 'p', 'div', 'ul', 'ol', 'table', 'pre', 'code']):
            self.process_element(element, chapter_num)
    
    def process_element(self, element, chapter_num=None):
        """Process individual HTML elements"""
        if element.name == 'h1':
            self.add_heading(element.get_text().strip(), 'ChapterTitle', chapter_num)
        elif element.name == 'h2':
            self.add_heading(element.get_text().strip(), 'CustomH2')
        elif element.name == 'h3':
            self.add_heading(element.get_text().strip(), 'CustomH3')
        elif element.name == 'h4':
            self.add_heading(element.get_text().strip(), 'CustomH4')
        elif element.name == 'p':
            self.add_paragraph(element)
        elif element.name in ['ul', 'ol']:
            self.add_list(element)
        elif element.name == 'table':
            self.add_table(element)
        elif element.name == 'pre':
            self.add_code_block(element)
        elif element.name == 'div':
            self.process_div(element)
    
    def add_heading(self, text, style_name, chapter_num=None):
        """Add a heading with specified style"""
        if chapter_num and style_name == 'ChapterTitle':
            # Add page break before new chapter (except first chapter)
            if chapter_num > 1:
                self.doc.add_page_break()
        
        heading = self.doc.add_paragraph(text, style=style_name)
        return heading
    
    def add_paragraph(self, element):
        """Add a paragraph with proper formatting"""
        text = element.get_text().strip()
        if not text:
            return
            
        # Check if paragraph has special classes
        classes = element.get('class', [])
        
        if 'highlight-box' in classes or 'feature-overview' in classes:
            # Add highlighted paragraph
            para = self.doc.add_paragraph()
            run = para.add_run(text)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            # Add background color (this is complex in python-docx, so we'll use bold instead)
            para.paragraph_format.left_indent = Inches(0.5)
            para.paragraph_format.right_indent = Inches(0.5)
        elif 'info-box' in classes or 'definition-box' in classes:
            # Add info box
            para = self.doc.add_paragraph()
            run = para.add_run(text)
            para.paragraph_format.left_indent = Inches(0.5)
        else:
            # Regular paragraph
            para = self.doc.add_paragraph(text)
    
    def add_list(self, element):
        """Add bulleted or numbered list"""
        for li in element.find_all('li'):
            text = li.get_text().strip()
            if text:
                if element.name == 'ul':
                    para = self.doc.add_paragraph(text, style='List Bullet')
                else:
                    para = self.doc.add_paragraph(text, style='List Number')
    
    def add_table(self, element):
        """Add table to document"""
        try:
            rows = element.find_all('tr')
            if not rows:
                return
                
            # Count maximum columns across all rows
            max_cols = 0
            for row in rows:
                cells = row.find_all(['th', 'td'])
                max_cols = max(max_cols, len(cells))
            
            if max_cols == 0:
                return
            
            # Create table
            table = self.doc.add_table(rows=len(rows), cols=max_cols)
            table.style = 'Table Grid'
            
            # Fill table data
            for row_idx, row in enumerate(rows):
                cells = row.find_all(['th', 'td'])
                for col_idx, cell in enumerate(cells):
                    if col_idx < max_cols and row_idx < len(table.rows):
                        try:
                            cell_text = cell.get_text().strip()
                            table.cell(row_idx, col_idx).text = cell_text
                            # Make header row bold
                            if row_idx == 0:
                                for paragraph in table.cell(row_idx, col_idx).paragraphs:
                                    for run in paragraph.runs:
                                        run.font.bold = True
                        except Exception as e:
                            print(f"Warning: Could not process table cell at row {row_idx}, col {col_idx}: {e}")
                            continue
        except Exception as e:
            print(f"Warning: Could not process table: {e}")
            # Add table content as text instead
            table_text = element.get_text().strip()
            if table_text:
                self.doc.add_paragraph(f"[TABLE CONTENT]: {table_text}")
    
    def add_code_block(self, element):
        """Add code block with proper formatting"""
        code_text = element.get_text().strip()
        if not code_text:
            return
            
        # Split code into lines and add each as a separate paragraph
        lines = code_text.split('\n')
        for line in lines:
            para = self.doc.add_paragraph(line, style='CodeBlock')
            # Add gray background (simplified approach)
            para.paragraph_format.left_indent = Inches(0.5)
    
    def process_div(self, element):
        """Process div elements based on their classes"""
        classes = element.get('class', [])
        
        # Skip navigation and header divs
        if any(cls in classes for cls in ['navigation', 'header', 'nav-button']):
            return
            
        # Process content divs
        if 'content-section' in classes or 'functionality-grid' in classes:
            for child in element.find_all(['h1', 'h2', 'h3', 'h4', 'p', 'ul', 'ol', 'table', 'pre']):
                self.process_element(child)
        elif 'function-card' in classes or 'benefit-item' in classes:
            # Process card content
            for child in element.find_all(['h4', 'p', 'ul']):
                self.process_element(child)
    
    def convert_files(self, html_files):
        """Convert multiple HTML files to a single DOCX"""
        # Add title page
        self.add_title_page()
        
        # Process table of contents first
        toc_file = None
        chapter_files = []
        
        for file_path in html_files:
            if 'table_of_contents' in file_path.lower():
                toc_file = file_path
            else:
                chapter_files.append(file_path)
        
        # Process table of contents
        if toc_file:
            print(f"Processing table of contents: {toc_file}")
            with open(toc_file, 'r', encoding='utf-8') as f:
                html_content = f.read()
            self.process_html_content(html_content)
            self.doc.add_page_break()
        
        # Sort chapter files numerically
        chapter_files.sort(key=lambda x: self.extract_chapter_number(x))
        
        # Process each chapter
        for idx, file_path in enumerate(chapter_files, 1):
            print(f"Processing chapter {idx}: {file_path}")
            with open(file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
            self.process_html_content(html_content, idx)
    
    def extract_chapter_number(self, filename):
        """Extract chapter number from filename"""
        match = re.search(r'chapter(\d+)', filename.lower())
        return int(match.group(1)) if match else 999
    
    def save_document(self, output_path):
        """Save the document to specified path"""
        self.doc.save(output_path)
        print(f"Document saved to: {output_path}")

def main():
    """Main function to run the conversion"""
    # Get current directory
    current_dir = Path('.')
    
    # Find all HTML files
    html_files = list(current_dir.glob('*.html'))
    
    if not html_files:
        print("No HTML files found in current directory!")
        return
    
    print(f"Found {len(html_files)} HTML files")
    
    # Create converter and process files
    converter = HTMLToDocxConverter()
    converter.convert_files([str(f) for f in html_files])
    
    # Save the document
    output_path = 'CRM_System_Documentation.docx'
    converter.save_document(output_path)
    
    print(f"\nConversion completed successfully!")
    print(f"Output file: {output_path}")

if __name__ == "__main__":
    main()
